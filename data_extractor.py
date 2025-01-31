from docx.opc.constants import RELATIONSHIP_TYPE as RT
from file_loader import PDFLoader, PPTLoader, DOCXLoader
from pdfminer.high_level import extract_text
import pdfplumber

class DataExtractor:
    """A class to extract text, links, images, and tables from various document formats."""

    def __init__(self, file_loader):
        """
        Initialize the DataExtractor with a specific file loader.

        Args:
            file_loader: An instance of PDFLoader, DOCXLoader, or PPTLoader that handles loading files.
        """
        self.file_loader = file_loader
        self.content = self.file_loader.load_file()  # Load the file using the provided file loader

    def extract_text(self):
        """
        Extract text content from the loaded file.

        Returns:
            str: The extracted text as a single string.
        """
        if isinstance(self.file_loader, PDFLoader):
            # Extract text from a PDF file using pdfminer
            text = extract_text(self.file_loader.file_path)
            if text.strip():
                return text  # Directly return PDF content if extracted text is not empty
        elif isinstance(self.file_loader, DOCXLoader):
            # Join all paragraph texts in a DOCX document
            return '\n'.join(paragraph.text for paragraph in self.content.paragraphs)
        elif isinstance(self.file_loader, PPTLoader):
            # Join all shape texts from each slide in a PPT presentation
            return '\n'.join(
                shape.text for slide in self.content.slides for shape in slide.shapes if hasattr(shape, "text")
            )
        return ""

    def extract_links(self):
        """
        Extract hyperlinks from the loaded file.

        Returns:
            list: A list of extracted hyperlinks.
        """
        if isinstance(self.file_loader, PDFLoader):
            return self.extract_pdf_links(self.file_loader.file_path)
        elif isinstance(self.file_loader, DOCXLoader):
            return self.extract_docx_links()
        elif isinstance(self.file_loader, PPTLoader):
            return self.extract_ppt_links()
        return []

    def extract_pdf_links(self, file_path):
        """
        Extract links from a PDF file.

        Args:
            file_path (str): Path to the PDF file.

        Returns:
            list: A list of extracted PDF links (URIs).
        """
        links = []
        with pdfplumber.open(file_path) as pdf:
            # Iterate through PDF pages to extract annotations containing URIs (links)
            for page in pdf.pages:
                if page.annots:
                    for annotation in page.annots:
                        if annotation.get("uri"):
                            links.append(annotation["uri"])
        return links

    def extract_docx_links(self):
        """
        Extract links from a DOCX file.

        Returns:
            list: A list of extracted hyperlinks in the DOCX document.
        """
        links = []
        for rel in self.content.part.rels.values():
            # Extract hyperlinks based on relationship type in the DOCX document
            if rel.reltype == RT.HYPERLINK:
                links.append(rel._target)
        return links

    def extract_ppt_links(self):
        """
        Extract links from a PPTX file.

        Returns:
            list: A list of extracted hyperlinks in the PPTX presentation.
        """
        links = []
        for slide in self.content.slides:
            for shape in slide.shapes:
                if shape.has_text_frame:
                    for paragraph in shape.text_frame.paragraphs:
                        for run in paragraph.runs:
                            # Check for hyperlinks in text runs and extract addresses
                            if run.hyperlink and run.hyperlink.address:
                                links.append(run.hyperlink.address)
        return links

    def extract_images(self):
        """
        Extract images from the loaded file.

        Returns:
            list: A list of extracted images (binary data).
        """
        if isinstance(self.file_loader, PDFLoader):
            return self.extract_pdf_images(self.file_loader.file_path)
        elif isinstance(self.file_loader, DOCXLoader):
            return self.extract_docx_images()
        elif isinstance(self.file_loader, PPTLoader):
            return self.extract_ppt_images()
        return []

    def extract_pdf_images(self, file_path):
        """
        Extract images from a PDF file.

        Args:
            file_path (str): Path to the PDF file.

        Returns:
            list: A list of images extracted from the PDF.
        """
        images = []
        with pdfplumber.open(file_path) as pdf:
            # Extract images from each PDF page
            for page in pdf.pages:
                if page.images:
                    images.extend(page.images)
        return images

    def extract_docx_images(self):
        """
        Extract images from a DOCX file.

        Returns:
            list: A list of images (binary data) in the DOCX document.
        """
        images = []
        for rel in self.content.part.rels.values():
            # Extract images based on relationships in the DOCX file
            if "image" in rel.target_ref:
                images.append(rel.target_part.blob)
        return images

    def extract_ppt_images(self):
        """
        Extract images from a PPTX file.

        Returns:
            list: A list of images (binary data) in the PPTX presentation.
        """
        images = []
        for slide in self.content.slides:
            for shape in slide.shapes:
                # Check if the shape is an image (Picture)
                if shape.shape_type == 13:  # Shape type 13 corresponds to Picture
                    images.append(shape.image.blob)
        return images

    def extract_tables(self):
        """
        Extract tables from the loaded file.

        Returns:
            list: A list of tables extracted from the file.
        """
        if isinstance(self.file_loader, PDFLoader):
            return self.extract_pdf_tables(self.file_loader.file_path)
        elif isinstance(self.file_loader, DOCXLoader):
            return self.extract_docx_tables()
        elif isinstance(self.file_loader, PPTLoader):
            return self.extract_ppt_tables()
        return []

    def extract_pdf_tables(self, file_path):
        """
        Extract tables from a PDF file.

        Args:
            file_path (str): Path to the PDF file.

        Returns:
            list: A list of tables extracted from the PDF.
        """
        tables = []
        with pdfplumber.open(file_path) as pdf:
            # Extract tables from each PDF page
            for page in pdf.pages:
                tables.extend(page.extract_tables())
        return tables

    def extract_docx_tables(self):
        """
        Extract tables from a DOCX file.

        Returns:
            list: A list of tables extracted from the DOCX document.
        """
        tables = []
        # Extract tables by reading rows and cells in DOCX tables
        for table in self.content.tables:
            table_data = [[cell.text for cell in row.cells] for row in table.rows]
            tables.append(table_data)
        return tables

    def extract_ppt_tables(self):
        """
        Extract tables from a PPTX file.

        Returns:
            list: A list of tables extracted from the PPTX presentation.
        """
        tables = []
        for slide in self.content.slides:
            for shape in slide.shapes:
                # Check if the shape contains a table and extract its content
                if shape.has_table:
                    table_data = [[cell.text for cell in row.cells] for row in shape.table.rows]
                    tables.append(table_data)
        return tables

    def extract_metadata(self):
        """
        Extract metadata from the loaded file.

        Returns:
            dict: A dictionary containing extracted metadata.
        """
        if isinstance(self.file_loader, PDFLoader):
            return self.extract_pdf_metadata(self.file_loader.file_path)
        elif isinstance(self.file_loader, DOCXLoader) or isinstance(self.file_loader, PPTLoader):
            return self.extract_document_metadata()
        return {}

    def extract_pdf_metadata(self, file_path):
        """
        Extract metadata from a PDF file.

        Args:
            file_path (str): Path to the PDF file.

        Returns:
            dict: A dictionary containing the PDF's metadata.
        """
        with pdfplumber.open(file_path) as pdf:
            return pdf.metadata

    def extract_document_metadata(self):
        """
        Extract metadata from a DOCX or PPTX file.

        Returns:
            dict: A dictionary containing the document's metadata.
        """
        core_properties = self.content.core_properties
        return {
            'title': core_properties.title,
            'author': core_properties.author,
            'subject': core_properties.subject,
            'keywords': core_properties.keywords,
            'created': core_properties.created,
            'modified': core_properties.modified,
        }
